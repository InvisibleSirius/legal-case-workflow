#!/usr/bin/env python3
"""Profile-aware structural checks for Chinese legal DOCX packages."""

from __future__ import annotations

import argparse
import json
import sys
import tempfile
import zipfile
from pathlib import Path
from typing import Any, Iterable
from xml.etree import ElementTree as ET

try:
    from docx import Document
    from docx.shared import Pt
except ImportError as exc:  # pragma: no cover - environment failure
    raise SystemExit("缺少 python-docx；请使用工作区随附的文档运行环境。") from exc


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
EP_NS = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
DEFAULT_FINAL_FORBIDDEN = ["【待确认】", "【待填写】", "【待核对】", "传票载", "未提供可供核验"]


def _paragraphs_in_table(table: Any) -> Iterable[Any]:
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _paragraphs_in_table(nested)


def _all_paragraphs(document: Any) -> Iterable[Any]:
    yield from document.paragraphs
    for table in document.tables:
        yield from _paragraphs_in_table(table)
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def _text_from_xml(raw: bytes) -> str:
    root = ET.fromstring(raw)
    return "".join(node.text or "" for node in root.iter() if node.tag == f"{{{W_NS}}}t")


def _app_pages(archive: zipfile.ZipFile) -> int | None:
    try:
        root = ET.fromstring(archive.read("docProps/app.xml"))
    except (KeyError, ET.ParseError):
        return None
    node = root.find(f"{{{EP_NS}}}Pages")
    if node is None or not node.text:
        return None
    try:
        return int(node.text)
    except ValueError:
        return None


def _explicit_east_asia_fonts(archive: zipfile.ZipFile) -> list[str]:
    fonts: set[str] = set()
    for name in archive.namelist():
        if not name.startswith("word/") or not name.endswith(".xml"):
            continue
        try:
            root = ET.fromstring(archive.read(name))
        except ET.ParseError:
            continue
        for node in root.iter(f"{{{W_NS}}}rFonts"):
            value = node.attrib.get(f"{{{W_NS}}}eastAsia")
            if value:
                fonts.add(value)
    return sorted(fonts)


def inspect_docx(path: Path) -> dict[str, Any]:
    info: dict[str, Any] = {"path": str(path), "zip_ok": False}
    with zipfile.ZipFile(path) as archive:
        bad_member = archive.testzip()
        if bad_member:
            raise ValueError(f"DOCX 压缩包损坏：{bad_member}")
        info["zip_ok"] = True
        text_parts: list[str] = []
        for name in archive.namelist():
            basename = Path(name).name
            if name.startswith("word/") and name.endswith(".xml") and basename.startswith(
                ("document", "header", "footer", "footnotes", "endnotes", "comments")
            ):
                try:
                    text_parts.append(_text_from_xml(archive.read(name)))
                except ET.ParseError:
                    pass
        info["visible_text"] = "\n".join(text_parts)
        info["app_pages"] = _app_pages(archive)
        info["media_count"] = sum(1 for name in archive.namelist() if name.startswith("word/media/"))
        info["explicit_east_asia_fonts"] = _explicit_east_asia_fonts(archive)

    document = Document(path)
    paragraphs = list(_all_paragraphs(document))
    sizes: list[float] = []
    paragraph_runs: list[dict[str, Any]] = []
    for paragraph in paragraphs:
        run_sizes: list[float] = []
        for run in paragraph.runs:
            if run.text.strip() and run.font.size is not None:
                value = round(run.font.size.pt, 2)
                sizes.append(value)
                run_sizes.append(value)
        paragraph_runs.append({"text": paragraph.text, "sizes_pt": run_sizes})
    info["direct_run_sizes_pt"] = sorted(set(sizes))
    info["paragraph_runs"] = paragraph_runs
    info["tables"] = [
        {"rows": len(table.rows), "columns": max((len(row.cells) for row in table.rows), default=0)}
        for table in document.tables
    ]
    return info


def _issue(kind: str, code: str, path: Path, message: str) -> dict[str, str]:
    return {"level": kind, "code": code, "file": str(path), "message": message}


def _contains_size(values: Iterable[float], expected: float, tolerance: float = 0.25) -> bool:
    return any(abs(value - expected) <= tolerance for value in values)


def _apply_rule(info: dict[str, Any], rule: dict[str, Any], path: Path) -> tuple[list[dict[str, str]], list[dict[str, str]]]:
    errors: list[dict[str, str]] = []
    warnings: list[dict[str, str]] = []
    tables = info.get("tables", [])

    expected_rows = rule.get("table_rows")
    expected_columns = rule.get("table_columns")
    if expected_rows is not None or expected_columns is not None:
        first = tables[0] if tables else None
        if first is None:
            errors.append(_issue("error", "missing_table", path, "画像要求主表格，但文档没有表格。"))
        else:
            if expected_rows is not None and first["rows"] != expected_rows:
                errors.append(_issue("error", "table_rows", path, f"主表格为 {first['rows']} 行，画像要求 {expected_rows} 行。"))
            if expected_columns is not None and first["columns"] != expected_columns:
                errors.append(_issue("error", "table_columns", path, f"主表格为 {first['columns']} 列，画像要求 {expected_columns} 列。"))

    max_pages = rule.get("max_app_pages")
    if max_pages is not None:
        pages = info.get("app_pages")
        if pages is None:
            warnings.append(_issue("warning", "pages_unknown", path, "文档属性未提供页数，仍需逐页渲染确认。"))
        elif pages > max_pages:
            errors.append(_issue("error", "page_count", path, f"文档属性页数为 {pages}，画像上限为 {max_pages}。"))

    forbidden_fonts = set(rule.get("forbidden_explicit_east_asia_fonts", []))
    used_forbidden = forbidden_fonts.intersection(info.get("explicit_east_asia_fonts", []))
    if used_forbidden:
        errors.append(_issue("error", "forbidden_font", path, f"显式使用画像禁止字体：{', '.join(sorted(used_forbidden))}。"))

    minimum_media = rule.get("required_media_min")
    if minimum_media is not None and info.get("media_count", 0) < minimum_media:
        errors.append(_issue("error", "missing_media", path, f"图片数量为 {info.get('media_count', 0)}，画像至少要求 {minimum_media} 个。"))

    visible_text = info.get("visible_text", "")
    for required in rule.get("required_text", []):
        if required not in visible_text:
            errors.append(_issue("error", "required_text", path, f"缺少画像要求文字：{required}"))

    required_size = rule.get("required_size_pt")
    if required_size is not None and not _contains_size(info.get("direct_run_sizes_pt", []), float(required_size)):
        errors.append(_issue("error", "required_size", path, f"未发现直接设置为 {required_size} 磅的文字。"))

    title_text = rule.get("title_text")
    title_size = rule.get("title_size_pt")
    if title_text and title_size is not None:
        title_paragraphs = [item for item in info.get("paragraph_runs", []) if title_text in item["text"]]
        if not title_paragraphs:
            errors.append(_issue("error", "missing_title", path, f"未找到标题：{title_text}"))
        elif not any(_contains_size(item["sizes_pt"], float(title_size)) for item in title_paragraphs):
            observed = sorted({size for item in title_paragraphs for size in item["sizes_pt"]})
            errors.append(_issue("error", "title_size", path, f"标题直接字号为 {observed or '未设置'}，画像要求 {title_size} 磅。"))

    return errors, warnings


def validate_package(package_dir: Path, profile: dict[str, Any], final: bool) -> dict[str, Any]:
    errors: list[dict[str, str]] = []
    warnings: list[dict[str, str]] = []
    files: list[dict[str, Any]] = []
    rules = profile.get("documents", [])
    matched_ids: set[str] = set()
    forbidden = profile.get("final_forbidden_phrases", DEFAULT_FINAL_FORBIDDEN)

    docx_paths = sorted(package_dir.rglob("*.docx"))
    for path in docx_paths:
        try:
            info = inspect_docx(path)
        except (OSError, ValueError, zipfile.BadZipFile) as exc:
            errors.append(_issue("error", "unreadable_docx", path, str(exc)))
            continue

        matched_rule = next(
            (
                rule
                for rule in rules
                if rule.get("filename_contains", "") in path.name
                and path.suffix.lower() == rule.get("extension", ".docx").lower()
            ),
            None,
        )
        if matched_rule:
            matched_ids.add(str(matched_rule.get("id", path.name)))
            rule_errors, rule_warnings = _apply_rule(info, matched_rule, path)
            errors.extend(rule_errors)
            warnings.extend(rule_warnings)

        if final:
            for phrase in forbidden:
                if phrase and phrase in info.get("visible_text", ""):
                    errors.append(_issue("error", "final_forbidden_phrase", path, f"正式版含内部过程语言：{phrase}"))

        files.append(
            {
                "path": str(path),
                "matched_profile_document": matched_rule.get("id") if matched_rule else None,
                "app_pages": info.get("app_pages"),
                "tables": info.get("tables"),
                "media_count": info.get("media_count"),
                "explicit_east_asia_fonts": info.get("explicit_east_asia_fonts"),
                "direct_run_sizes_pt": info.get("direct_run_sizes_pt"),
            }
        )

    if not docx_paths:
        errors.append(_issue("error", "no_docx", package_dir, "交付目录中没有 DOCX 文件。"))

    for path in sorted(package_dir.rglob("*.doc")):
        warnings.append(_issue("warning", "legacy_doc_unchecked", path, "旧版 .doc 未做结构检查，请先转换并逐页比对。"))

    for rule in rules:
        rule_id = str(rule.get("id", "unknown"))
        if rule_id not in matched_ids:
            warnings.append(
                _issue(
                    "warning",
                    "profile_document_absent",
                    package_dir,
                    f"画像中的 {rule_id} 未出现在本次文书包；如采用 requested-only 可忽略。",
                )
            )

    return {
        "status": "fail" if errors else "pass",
        "package_dir": str(package_dir),
        "profile_id": profile.get("profile_id"),
        "final_mode": final,
        "errors": errors,
        "warnings": warnings,
        "files": files,
    }


def _self_test() -> int:
    with tempfile.TemporaryDirectory(prefix="legal-docx-validator-") as temp_dir:
        root = Path(temp_dir)
        path = root / "授权委托书.docx"
        document = Document()
        title = document.add_paragraph()
        title_run = title.add_run("授权委托书")
        title_run.font.size = Pt(22)
        body = document.add_paragraph()
        body_run = body.add_run("广东勤通律师事务所")
        body_run.font.size = Pt(15)
        document.save(path)
        profile = {
            "profile_id": "self-test",
            "documents": [
                {
                    "id": "authorization",
                    "filename_contains": "授权委托书",
                    "extension": ".docx",
                    "title_text": "授权委托书",
                    "title_size_pt": 22,
                    "required_size_pt": 15,
                }
            ],
        }
        report = validate_package(root, profile, final=True)
        print(json.dumps(report, ensure_ascii=False, indent=2))
        return 0 if report["status"] == "pass" else 1


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="按版本画像检查法律文书 DOCX 包。")
    parser.add_argument("--package-dir", type=Path, help="待检查的交付目录")
    parser.add_argument("--profile", type=Path, help="JSON 格式版本画像")
    parser.add_argument("--final", action="store_true", help="启用正式终稿禁用语检查")
    parser.add_argument("--report", type=Path, help="另存 JSON 报告")
    parser.add_argument("--self-test", action="store_true", help="运行内置测试")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    if args.self_test:
        return _self_test()
    if args.package_dir is None or args.profile is None:
        raise SystemExit("必须同时提供 --package-dir 和 --profile，或使用 --self-test。")
    if not args.package_dir.is_dir():
        raise SystemExit(f"交付目录不存在：{args.package_dir}")
    if not args.profile.is_file():
        raise SystemExit(f"画像文件不存在：{args.profile}")
    profile = json.loads(args.profile.read_text(encoding="utf-8"))
    report = validate_package(args.package_dir, profile, args.final)
    rendered = json.dumps(report, ensure_ascii=False, indent=2)
    print(rendered)
    if args.report:
        args.report.parent.mkdir(parents=True, exist_ok=True)
        args.report.write_text(rendered + "\n", encoding="utf-8")
    return 0 if report["status"] == "pass" else 1


if __name__ == "__main__":
    sys.exit(main())
